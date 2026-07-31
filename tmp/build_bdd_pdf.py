from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CARD_TITLE = "MCSL fedex rest - fetch negotiate rate"
OUTPUT = Path("/Users/ashokkumarn/Documents/Pluginhive/AILearning/MCSLDomainExpert/out/mcsl_fedex_rest_fetch_negotiate_rate_bdd.docx")


SCENARIOS = [
    (
        "Domestic shipment excludes FedEx One Rate when the exclude-One-Rate behavior is active",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is domestic",
            "And the package uses a custom box with valid dimensions",
            "When the customer requests shipping rates at checkout",
            "Then the FedEx rate request should exclude One Rate pricing",
            "And the response should not contain FedEx One Rate charges",
            "And only the applicable negotiated or selected standard FedEx rates should be returned",
            "And the rate shown at checkout should match the rate response",
            "And the order summary shipping rate should match the checkout rate",
        ],
    ),
    (
        "Domestic shipment returns normal domestic rates when One Rate exclusion is not active",
        [
            "Given the FedEx REST rate changes feature toggle is OFF",
            "And the shipment is domestic",
            "And the package uses a custom box with valid dimensions",
            "When the customer requests shipping rates at checkout",
            "Then the legacy FedEx rating flow should be used",
            "And the system should behave as before the new change",
            "And the checkout rate should match the returned FedEx response",
            "And the order summary shipping rate should match the checkout rate",
        ],
    ),
    (
        "International shipment uses old API when feature toggle is OFF",
        [
            "Given the FedEx REST rate changes feature toggle is OFF",
            "And the shipment is international",
            "When the customer requests shipping rates at checkout",
            "Then the legacy rates API should be called",
            "And duties and taxes data should not be sent in the request",
            "And the returned rate should match the rate shown at checkout",
            "And the order summary should match the checkout rate",
        ],
    ),
    (
        "International shipment uses new transit API with duties and taxes disabled",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is international",
            "And the merchant has disabled estimated duties and taxes",
            "When the customer requests shipping rates at checkout",
            "Then the new FedEx transit API should be called",
            "And edtRequestType should be sent as NONE",
            "And the checkout rate should be returned successfully without duties and taxes added",
            "And the order summary should match the checkout rate",
        ],
    ),
    (
        "International shipment uses new transit API with duties and taxes enabled",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is international",
            "And the merchant has enabled estimated duties and taxes",
            "When the customer requests shipping rates at checkout",
            "Then the new FedEx transit API should be called",
            "And edtRequestType should be sent as ALL",
            "And duties and taxes should be included in the returned rating data",
            "And the shipping charge shown at checkout should reflect the returned duties and taxes behavior correctly",
            "And the order summary should match the checkout rate",
        ],
    ),
    (
        "Domestic shipment must not send international duties and taxes fields",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is domestic",
            "And the merchant has enabled estimated duties and taxes",
            "When the customer requests shipping rates at checkout",
            "Then international duties and taxes request fields should not be sent",
            "And the domestic rate response should still be returned successfully",
            "And checkout and order summary should remain consistent",
        ],
    ),
    (
        "Rate request type LIST shows only list rate service at checkout",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the carrier rate request type is set to LIST",
            "And the shipment is valid for FedEx rating",
            "When the customer requests shipping rates at checkout",
            "Then the FedEx request should ask for both LIST and ACCOUNT rates",
            "And the checkout should display only the LIST rate result for the service",
            "And the order summary should match the checkout rate",
        ],
    ),
    (
        "Rate request type ACCOUNT shows only negotiated account rate at checkout",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the carrier rate request type is set to ACCOUNT",
            "And the shipment is valid for FedEx rating",
            "When the customer requests shipping rates at checkout",
            "Then the FedEx request should ask for both LIST and ACCOUNT rates",
            "And the checkout should display only the ACCOUNT negotiated rate for the service",
            "And the order summary should match the checkout rate",
        ],
    ),
    (
        "Rate request type PREFERRED returns preferred-priced services",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the carrier rate request type is set to PREFERRED",
            "And the shipment is valid for FedEx rating",
            "When the customer requests shipping rates at checkout",
            "Then the FedEx request should use PREFERRED pricing selection",
            "And the checkout should display only the preferred-priced service result",
            "And the displayed rate should match the FedEx response selected by the system",
        ],
    ),
    (
        "One Rate inclusion setting ON returns lower-rate selection behavior",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is domestic",
            "And the merchant has enabled the One Rate related carrier setting",
            "When the customer requests shipping rates at checkout",
            "Then the rate display option sent to FedEx should be LOWER_RATE",
            "And the returned service price should follow the lower-rate selection logic",
            "And checkout and order summary should match the selected returned rate",
        ],
    ),
    (
        "One Rate inclusion setting OFF excludes F1R rates",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is domestic",
            "And the merchant has disabled the One Rate related carrier setting",
            "When the customer requests shipping rates at checkout",
            "Then the rate display option sent to FedEx should be SELECTED_RATES_EXCLUDING_F1R",
            "And FedEx One Rate results should not be returned",
            "And checkout and order summary should match the selected returned rate",
        ],
    ),
    (
        "Carrier packaging with One Rate exclusion does not break rate retrieval",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is domestic",
            "And the package uses FedEx carrier packaging",
            "When the customer requests shipping rates at checkout",
            "Then the system should return valid rates without error",
            "And excluded One Rate behavior should still be honored as configured",
            "And checkout and order summary should match",
        ],
    ),
    (
        "Custom packaging with One Rate exclusion does not break rate retrieval",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is domestic",
            "And the package uses custom packaging",
            "When the customer requests shipping rates at checkout",
            "Then the system should return valid rates without error",
            "And excluded One Rate behavior should still be honored as configured",
            "And checkout and order summary should match",
        ],
    ),
    (
        "Multiple-package international shipment is rated successfully with the new API",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is international",
            "And the shipment contains multiple packages",
            "When the customer requests shipping rates at checkout",
            "Then the new international rating flow should return valid rates",
            "And duties and taxes behavior should follow the merchant setting",
            "And checkout and order summary should match the returned result",
        ],
    ),
    (
        "International freight shipment is rated successfully with the new API",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is international freight",
            "When the customer requests shipping rates at checkout",
            "Then the new international rating flow should return valid rates without breaking",
            "And the selected rate should be displayed correctly at checkout",
            "And the order summary should match the checkout rate",
        ],
    ),
    (
        "International Saturday delivery rating does not break the new flow",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is international",
            "And Saturday delivery is applicable in the request",
            "When the customer requests shipping rates at checkout",
            "Then the new international rating flow should return a valid result or a clean unsupported response",
            "And the system should not show mismatched or duplicated rates",
            "And the order summary should remain consistent with checkout",
        ],
    ),
    (
        "Unsupported One Rate or SmartPost combinations for international shipments do not break checkout",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the shipment is international",
            "And the request contains a service combination unsupported for international shipment such as One Rate or SmartPost",
            "When the customer requests shipping rates at checkout",
            "Then the system should not fail unexpectedly",
            "And unsupported services should not be displayed as purchasable checkout options",
            "And valid supported services should still be shown when available",
        ],
    ),
    (
        "No mismatch between checkout rate and order summary after rate selection",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And a FedEx rate is displayed at checkout",
            "When the customer completes the order",
            "Then the shipping amount stored in the order summary should exactly match the checkout rate",
            "And there should be no mismatch between API response, checkout display, and order summary",
        ],
    ),
    (
        "No negotiated rate available for account-rated request",
        [
            "Given the FedEx REST rate changes feature toggle is ON",
            "And the carrier rate request type is set to ACCOUNT",
            "And FedEx does not return an account-specific negotiated rate for the shipment",
            "When the customer requests shipping rates at checkout",
            "Then the system should handle the response gracefully",
            "And it should not display an incorrect One Rate result as a fallback",
            "And it should either show a valid fallback per business rule or no rate with a clear outcome",
        ],
    ),
    (
        "Toggle-disabled store does not use the new international rate logic",
        [
            "Given the store does not satisfy the FedEx REST rate changes toggle condition",
            "And the shipment is international",
            "When the customer requests shipping rates at checkout",
            "Then the new transit API should not be used",
            "And the store should continue using the old rating behavior",
            "And the response should remain stable without regression",
        ],
    ),
]


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "MCSL FedEx REST BDD Scenarios"
    header.style = doc.styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header.runs:
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    add_page_number(footer)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BDD Test Scenarios")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(CARD_TITLE)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    feature = doc.add_paragraph()
    feature.add_run("Feature: ").bold = True
    feature.add_run("FedEx REST rates exclude One Rate and return the correct negotiated or selected rate")

    doc.add_paragraph("Background:", style="Heading 2")
    background = [
        "Given the merchant has FedEx REST configured in MCSL",
        "And the store is eligible for the FedEx REST rate changes feature",
        "And the product, address, and package data are valid for FedEx rating",
        "And checkout rate display, order rate summary, and API response can be verified together",
    ]
    for line in background:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)

    doc.add_paragraph("Scenarios", style="Heading 1")
    for idx, (title_text, lines) in enumerate(SCENARIOS, start=1):
        h = doc.add_paragraph(style="Heading 2")
        h.add_run(f"Scenario {idx}: {title_text}")
        for line in lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(line)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
