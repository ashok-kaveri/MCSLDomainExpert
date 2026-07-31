from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE = "Support Guide"
SUBTITLE = "FedEx One Rate: One rate inclusion logic for domestic and international"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_number_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def build_doc(output_path: Path):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    set_page_number_footer(section)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(TITLE)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 52, 91)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sub_run = subtitle.add_run(SUBTITLE)
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(89, 98, 117)

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    meta.columns[0].width = Inches(1.8)
    meta.columns[1].width = Inches(5.8)
    metadata = [
        ("Card", "Trello card 6jBLYcry"),
        ("Ticket", "Zendesk #382425"),
        ("Area", "FedEx rating / label-generation settings"),
        ("Audience", "Support and QA teams"),
    ]
    for row, (label, value) in zip(meta.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "E9EEF7")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10.5)
    document.add_paragraph()

    headings = [
        "Issue Summary",
        "Customer-Facing Symptoms",
        "Support Understanding",
        "Recommended Support Workflow",
        "Expected Fixed Behaviour",
        "QA Test Scenarios",
        "Support Notes",
    ]
    for style_name in ("Heading 1", "Heading 2"):
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(31, 52, 91)

    document.add_heading("Issue Summary", level=1)
    p = document.add_paragraph()
    p.add_run("Problem: ").bold = True
    p.add_run(
        "The FedEx One Rate setting and Duties & Taxes setting are currently tied to one shared toggle path. "
        "Because international rate requests are routed through the Transit API, enabling this shared behavior can cause "
        "international rate failures on stores where the toggle is enabled."
    )
    p = document.add_paragraph()
    p.add_run("Goal of the change: ").bold = True
    p.add_run(
        "Split One Rate and Duties & Taxes into two independent settings so domestic One Rate behavior and "
        "international duties handling do not overlap."
    )

    document.add_heading("Customer-Facing Symptoms", level=1)
    add_bullets(
        document,
        [
            "Domestic FedEx One Rate eligible flows may work, but international rate requests fail when the shared toggle is enabled.",
            "Support may see stores reporting that international FedEx rates stop returning after enabling One Rate-related behavior.",
            "Merchants can interpret the problem as a general FedEx outage even though the root cause is settings overlap.",
        ],
    )

    document.add_heading("Support Understanding", level=1)
    add_bullets(
        document,
        [
            "One Rate should affect only eligible domestic FedEx flows.",
            "Duties & Taxes should affect only international rating / label-generation behavior where customs handling is relevant.",
            "Transit API international requests must not inherit One Rate logic accidentally.",
            "The two settings must save, persist, and operate independently.",
        ],
    )

    document.add_heading("Recommended Support Workflow", level=1)
    add_numbered(
        document,
        [
            "Confirm whether the failing shipment is domestic or international.",
            "Check the FedEx carrier settings and capture the current values for One Rate and Duties & Taxes separately.",
            "If the issue is an international failure, verify whether One Rate is enabled and whether the failure started after that change.",
            "Reproduce with one domestic order and one international order to confirm the settings are behaving independently.",
            "Validate whether rates return successfully before and after toggling each setting individually.",
            "Share request / error evidence with engineering if the international Transit API call still fails when Duties & Taxes is disabled or when only One Rate is enabled.",
        ],
    )

    document.add_heading("Expected Fixed Behaviour", level=1)
    add_bullets(
        document,
        [
            "Domestic FedEx One Rate eligible orders continue to return One Rate services correctly.",
            "International FedEx Transit API requests no longer fail because One Rate is enabled.",
            "Duties & Taxes settings apply only to the intended international customs path.",
            "Turning one setting on or off does not mutate or override the other setting.",
            "Reopening FedEx carrier settings shows both saved values persisted independently.",
        ],
    )

    document.add_heading("QA Test Scenarios", level=1)
    scenarios = [
        (
            "TC-1: Domestic FedEx One Rate works independently when One Rate is enabled and Duties & Taxes is disabled",
            [
                "Given a store with FedEx configured and a domestic US order eligible for FedEx One Rate",
                "When the merchant enables only the One Rate setting and requests rates or generates a label",
                "Then One Rate services should be returned correctly for the domestic flow",
                "And the request should not depend on Duties & Taxes settings",
            ],
        ),
        (
            "TC-2: International rate request succeeds when One Rate is enabled but Duties & Taxes is disabled",
            [
                "Given a store with FedEx configured and an international order routed through the Transit API",
                "When the merchant enables One Rate and keeps Duties & Taxes disabled",
                "Then the international rate request should not fail because of the One Rate toggle",
                "And One Rate logic should not break the international Transit API flow",
            ],
        ),
        (
            "TC-3: International rate request succeeds when Duties & Taxes is enabled and One Rate is disabled",
            [
                "Given a store with FedEx configured and an international order that requires duties and customs handling",
                "When the merchant enables only Duties & Taxes and requests rates or generates a label",
                "Then the international request should succeed",
                "And duties and taxes behavior should apply without depending on One Rate",
            ],
        ),
        (
            "TC-4: Both settings can be enabled without overlapping behavior",
            [
                "Given a store with FedEx configured for both domestic and international shipment testing",
                "When the merchant enables both One Rate and Duties & Taxes",
                "Then domestic eligible orders should still use One Rate behavior only where applicable",
                "And international orders should continue to use duties and taxes handling without Transit API failure",
            ],
        ),
    ]
    for title_text, lines in scenarios:
        document.add_heading(title_text, level=2)
        add_bullets(document, lines)

    document.add_heading("Support Notes", level=1)
    notes_box = document.add_table(rows=1, cols=1)
    notes_box.alignment = WD_TABLE_ALIGNMENT.LEFT
    notes_box.style = "Table Grid"
    note_cell = notes_box.cell(0, 0)
    set_cell_shading(note_cell, "F7F9FC")
    note_cell.text = (
        "Important: if a merchant reports only international rate failures after enabling a FedEx setting, "
        "verify that One Rate and Duties & Taxes are no longer sharing one implementation path. "
        "This card is specifically about separating those controls to avoid Transit API regressions."
    )
    for paragraph in note_cell.paragraphs:
      for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)

    document.save(output_path)


if __name__ == "__main__":
    output_dir = Path("/Users/ashokkumarn/Documents/Pluginhive/AILearning/MCSLDomainExpert/output/support_guides")
    output_dir.mkdir(parents=True, exist_ok=True)
    build_doc(output_dir / "fedex_one_rate_support_guide.docx")
